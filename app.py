import requests
from bs4 import BeautifulSoup
from docx import Document
from io import BytesIO
import streamlit as st
from urllib.parse import urlparse
import re

# ScraperAPI base URL and API key
api_url = "http://api.scraperapi.com"
api_key = st.secrets["SCRAPER_API_KEY"]  # Replace with your ScraperAPI key

# Function to scrape using ScraperAPI
def scrape_with_scraperapi(url):
    params = {
        'api_key': api_key,
        'url': url,
    }
    try:
        response = requests.get(api_url, params=params, timeout=5)  # Reduced timeout
        response.raise_for_status()  # Raise an exception for unsuccessful requests
    except requests.exceptions.RequestException as e:
        st.error(f"Error fetching webpage: {e}")
        return None

    # Use the faster 'lxml' parser
    return BeautifulSoup(response.content, 'lxml')


# Optimized function to scrape and save content to DOCX using ScraperAPI
def scrape_and_save_to_docx(url):
    soup = scrape_with_scraperapi(url)
    if soup is None:
        return None

    doc = Document()

    # Extract and clean domain name for DOCX heading and file naming
    domain = urlparse(url).netloc or url
    valid_filename = re.sub(r'[^a-zA-Z0-9_-]', '_', domain)

    doc.add_heading(f'Content from {domain}', level=1)

    # Handle tags efficiently
    for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'ol', 'img', 'audio', 'video']):
        if tag.name in ['h1', 'h2', 'h3', 'h4']:
            doc.add_heading(tag.get_text(strip=True), level=int(tag.name[1]))
        elif tag.name == 'p':
            doc.add_paragraph(tag.get_text(strip=True))
        elif tag.name == 'ul':
            for li in tag.find_all('li'):
                doc.add_paragraph(f"• {li.get_text(strip=True)}", style='List Bullet')
        elif tag.name == 'ol':
            for i, li in enumerate(tag.find_all('li')):
                doc.add_paragraph(f"{i+1}. {li.get_text(strip=True)}", style='List Number')
        elif tag.name == 'img':
            doc.add_paragraph(f"[Image Placeholder: {tag.get('src', 'Image URL')}]")
        elif tag.name == 'audio':
            doc.add_paragraph(f"[Audio Placeholder: {tag.get('src', 'Audio URL')}]")
        elif tag.name == 'video':
            doc.add_paragraph(f"[Video Placeholder: {tag.get('src', 'Video URL')}]")

    # Save to DOCX in memory (BytesIO)
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    return doc_io, f"{valid_filename}.docx"


# Function to display DOCX content on the screen
def display_docx_content(docx_io):
    document = Document(docx_io)
    return "\n".join([para.text for para in document.paragraphs])


# Streamlit app
st.title("Rapid Web Scraper to DOCX using ScraperAPI")

# Input for URL
url = st.text_input("Enter the URL of the webpage to scrape")

# Button to scrape and generate DOCX file
if st.button("Generate and View DOCX"):
    if url:
        with st.spinner("Scraping content and generating DOCX..."):
            docx_io, file_name = scrape_and_save_to_docx(url)

            if docx_io:
                st.success(f"DOCX file '{file_name}' generated successfully!")

                # Display the DOCX content on the screen
                docx_content = display_docx_content(docx_io)
                st.text_area("DOCX Content", docx_content, height=300)

                # Download button for the DOCX file
                st.download_button(
                    label="Download DOCX",
                    data=docx_io,
                    file_name=file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    else:
        st.error("Please enter a valid URL.")
